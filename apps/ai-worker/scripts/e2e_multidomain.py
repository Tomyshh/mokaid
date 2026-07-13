"""End-to-end multi-domain pipeline validation with REAL API keys.

Exercises the exact production code path of the ingestion + retrieval +
generation pipeline on four business domains with mixed input formats:

  1. Finance      — XLSX accounts + text instruction
  2. Legal        — contract PDF + DOCX annex
  3. Medical      — plain-text notes + PDF report
  4. SEO          — pure text brief (deliverable generation)

For each scenario: extract (app.memory.extractors) -> chunk + contextualize
(app.memory.ingestion) -> embed (real OpenAI embeddings) -> retrieve by
cosine similarity -> generate an answer with the retrieved context (real
LLM) -> assert the answer contains the ground-truth facts.

Cost: a handful of embedding batches + 4 short completions (< $0.05).

Run from apps/ai-worker (reads .env):  .venv/bin/python scripts/e2e_multidomain.py
"""

import asyncio
import io
import math
import sys
from dataclasses import dataclass

sys.path.insert(0, ".")

from app import llm  # noqa: E402
from app.memory import extractors  # noqa: E402
from app.memory.ingestion import chunk_text, contextualize  # noqa: E402

# ---------- fixtures ----------


def make_finance_xlsx() -> bytes:
    import openpyxl

    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Comptes 2026"
    rows = [
        ["Poste", "T1", "T2", "T3"],
        ["Revenus", 120000, 135000, 158500],
        ["Charges salariales", 60000, 62000, 64000],
        ["Marketing", 8000, 9500, 12000],
        ["Résultat net", 52000, 63500, 82500],
    ]
    for row in rows:
        sheet.append(row)
    buf = io.BytesIO()
    workbook.save(buf)
    return buf.getvalue()


def make_legal_pdf() -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "CONTRAT DE PRESTATION DE SERVICES\n\n"
        "Article 7 - Penalites de retard\n"
        "En cas de retard de livraison, le prestataire versera une penalite\n"
        "de 250 euros par jour ouvrable de retard, plafonnee a 10% du\n"
        "montant total du contrat.\n\n"
        "Article 8 - Resiliation\n"
        "Chaque partie peut resilier avec un preavis de 60 jours.",
    )
    data = doc.tobytes()
    doc.close()
    return data


def make_legal_docx() -> bytes:
    import docx

    document = docx.Document()
    document.add_heading("Annexe A — Livrables", level=1)
    document.add_paragraph(
        "Le prestataire livrera un audit de sécurité complet ainsi qu'un "
        "rapport de conformité RGPD avant le 15 septembre 2026."
    )
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def make_medical_pdf() -> bytes:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "COMPTE RENDU D'HOSPITALISATION\n\n"
        "Patient admis pour douleurs thoraciques. ECG normal.\n"
        "Troponine a 0,02 ng/mL (normale). Diagnostic retenu:\n"
        "douleur parietale, pas de syndrome coronarien aigu.\n"
        "Sortie avec traitement antalgique simple.",
    )
    data = doc.tobytes()
    doc.close()
    return data


MEDICAL_NOTES = (
    "Suivi du patient Dupont : tension artérielle stabilisée à 128/82 sous "
    "périndopril 5 mg. Prochain contrôle biologique (créatinine, kaliémie) "
    "dans 3 mois."
)

SEO_BRIEF = (
    "Brief SEO : rédiger un plan de contenu pour un cabinet d'expertise "
    "comptable à Lyon ciblant les mots-clés 'expert comptable lyon' et "
    "'création SASU'. Inclure 5 idées d'articles avec leurs H1."
)


# ---------- pipeline helpers (production code paths) ----------


@dataclass
class Corpus:
    chunks: list[str]
    vectors: list[list[float]]


def cosine(a: list[float], b: list[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b, strict=True))
    norm = math.sqrt(sum(x * x for x in a)) * math.sqrt(sum(y * y for y in b))
    return dot / norm if norm else 0.0


async def index_documents(docs: list[tuple[str, bytes | str, str]]) -> Corpus:
    """(title, payload, filename) -> extracted + chunked + embedded corpus."""
    all_chunks: list[str] = []
    for title, payload, filename in docs:
        if isinstance(payload, bytes):
            result = extractors.extract_bytes(payload, filename=filename)
            assert result is not None, f"extraction failed for {filename}"
            text = result.text
            print(f"    extracted {filename}: format={result.format} chars={len(text)}")
        else:
            text = payload
        chunks = contextualize(chunk_text(text), title)
        all_chunks.extend(chunks)
    vectors = await llm.embed(all_chunks)
    print(f"    indexed {len(all_chunks)} chunks ({len(vectors[0])} dims)")
    return Corpus(all_chunks, vectors)


async def retrieve(corpus: Corpus, query: str, k: int = 3) -> list[str]:
    [qvec] = await llm.embed([query])
    ranked = sorted(
        zip(corpus.chunks, corpus.vectors, strict=True),
        key=lambda pair: cosine(qvec, pair[1]),
        reverse=True,
    )
    return [chunk for chunk, _ in ranked[:k]]


async def answer(query: str, context_chunks: list[str]) -> str:
    context = "\n\n---\n\n".join(context_chunks)
    return await llm.chat(
        system=(
            "Tu réponds uniquement à partir du contexte fourni. "
            "Sois précis et factuel, cite les chiffres exacts."
        ),
        user=f"Contexte:\n{context}\n\nQuestion: {query}",
        max_tokens=400,
        quality="fast",
    )


def check(label: str, condition: bool, detail: str = "") -> bool:
    mark = "PASS" if condition else "FAIL"
    print(f"    [{mark}] {label}" + (f" — {detail}" if detail and not condition else ""))
    return condition


# ---------- scenarios ----------


async def scenario_finance() -> bool:
    print("\n[1/4] Finance — XLSX de comptes + question texte")
    corpus = await index_documents([("Comptes 2026", make_finance_xlsx(), "comptes.xlsx")])
    query = "Quel est le montant des revenus au troisième trimestre (T3) ?"
    top = await retrieve(corpus, query)
    ok_retrieval = check("retrieval: chunk des revenus au top", any("158500" in c for c in top))
    reply = await answer(query, top)
    print(f"    réponse: {reply[:160]}...")
    ok_answer = check("réponse: contient 158 500", "158 500" in reply or "158500" in reply or "158,500" in reply)
    return ok_retrieval and ok_answer


async def scenario_legal() -> bool:
    print("\n[2/4] Juridique — contrat PDF + annexe DOCX")
    corpus = await index_documents(
        [
            ("Contrat de prestation", make_legal_pdf(), "contrat.pdf"),
            ("Annexe A", make_legal_docx(), "annexe.docx"),
        ]
    )
    query_pdf = "Quelle est la pénalité en cas de retard de livraison ?"
    top_pdf = await retrieve(corpus, query_pdf)
    ok_pdf = check("retrieval PDF: clause pénalités trouvée", any("250" in c for c in top_pdf))

    query_docx = "Quels livrables sont prévus dans l'annexe ?"
    top_docx = await retrieve(corpus, query_docx)
    ok_docx = check(
        "retrieval DOCX: livrables de l'annexe trouvés",
        any("RGPD" in c or "audit" in c.lower() for c in top_docx),
    )

    reply = await answer(query_pdf, top_pdf)
    print(f"    réponse: {reply[:160]}...")
    ok_answer = check("réponse: cite 250 euros/jour", "250" in reply)
    return ok_pdf and ok_docx and ok_answer


async def scenario_medical() -> bool:
    print("\n[3/4] Médical — notes texte + compte rendu PDF")
    corpus = await index_documents(
        [
            ("Notes de suivi Dupont", MEDICAL_NOTES, "notes.txt"),
            ("Compte rendu d'hospitalisation", make_medical_pdf(), "cr.pdf"),
        ]
    )
    query = "Quel diagnostic a été retenu lors de l'hospitalisation ?"
    top = await retrieve(corpus, query)
    ok_retrieval = check(
        "retrieval: diagnostic dans le contexte",
        any("parietale" in c or "pariétale" in c for c in top),
    )
    reply = await answer(query, top)
    print(f"    réponse: {reply[:160]}...")
    ok_answer = check(
        "réponse: mentionne douleur pariétale / absence de SCA",
        "pariétale" in reply.lower() or "parietale" in reply.lower() or "coronarien" in reply.lower(),
    )

    query2 = "Quel est le traitement antihypertenseur du patient Dupont ?"
    top2 = await retrieve(corpus, query2)
    ok_cross = check("retrieval croisé: périndopril depuis les notes", any("périndopril" in c for c in top2))
    return ok_retrieval and ok_answer and ok_cross


async def scenario_seo() -> bool:
    print("\n[4/4] SEO — brief texte pur, production d'un livrable")
    deliverable = await llm.chat(
        system=(
            "Tu es un consultant SEO senior. Produis un livrable markdown "
            "structuré et actionnable, en français."
        ),
        user=SEO_BRIEF,
        max_tokens=900,
        quality="fast",
    )
    print(f"    livrable: {len(deliverable)} caractères")
    ok_len = check("livrable: taille suffisante", len(deliverable) > 500)
    ok_kw = check(
        "livrable: reprend les mots-clés du brief",
        "expert comptable" in deliverable.lower() and "sasu" in deliverable.lower(),
    )
    ok_structure = check("livrable: structuré (titres markdown)", deliverable.count("#") >= 3)

    # Le livrable doit lui-même être ré-ingérable (boucle de knowledge).
    chunks = contextualize(chunk_text(deliverable), "Plan de contenu SEO Lyon")
    ok_loop = check("boucle knowledge: livrable chunkable", len(chunks) >= 1)
    return ok_len and ok_kw and ok_structure and ok_loop


async def main() -> int:
    if not llm.is_configured():
        print("ERREUR: aucune clé API configurée (.env) — scénarios e2e impossibles.")
        return 2

    results = {
        "finance": await scenario_finance(),
        "juridique": await scenario_legal(),
        "médical": await scenario_medical(),
        "seo": await scenario_seo(),
    }

    print("\n" + "=" * 50)
    for name, ok in results.items():
        print(f"  {name:12s} {'PASS' if ok else 'FAIL'}")
    all_ok = all(results.values())
    print(f"\n{'TOUS LES SCÉNARIOS PASSENT' if all_ok else 'DES SCÉNARIOS ONT ÉCHOUÉ'}")
    return 0 if all_ok else 1


if __name__ == "__main__":
    raise SystemExit(asyncio.run(main()))
